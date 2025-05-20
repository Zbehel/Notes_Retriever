import pypdf
from docx import Document
from transformers.pipelines import pipeline
from sentence_transformers import SentenceTransformer
import streamlit as st
import numpy as np
import os

emb_model = "intfloat/multilingual-e5-large-instruct"
emb_model2 = "DeepPavlov/distilrubert-small-cased-conversational"
def extract_text(file):
    text = ""
    # Check if the input is a file path (string) or a file-like object
    if isinstance(file, str):
        file_name = os.path.basename(file)
        try:
            with open(file, 'rb') as f: # Open in binary mode
                if file_name.endswith(".pdf"):
                    print('Processing pdf file.................\n')
                    reader = pypdf.PdfReader(f)
                    for page in reader.pages:
                        text += page.extract_text() + "\\n"
                elif file_name.endswith(".docx"):
                    document = Document(f)
                    print('Processing DOCX file.................\n')
                    for paragraph in document.paragraphs:
                        if paragraph.text.strip():  # Check if the paragraph is not empty
                            text += paragraph.text + "\\n"
        except FileNotFoundError:
            st.error(f"Error: File not found at {file}")
            return ""
        except Exception as e:
            st.error(f"Error reading {file_name}: {e}")
            return ""
    else: # Assume it's a file-like object (e.g., from Streamlit file_uploader)
        file_name = file.name
        try:
            if file_name.endswith(".pdf"):
                reader = pypdf.PdfReader(file)
                for page in reader.pages:
                    text += page.extract_text() + "\\n"
            elif file_name.endswith(".docx"):
                document = Document(file)
                for paragraph in document.paragraphs:
                    text += paragraph.text + "\\n"
        except Exception as e:
            st.error(f"Error reading {file_name}: {e}")
            return ""
    return text

def chunk_text(text, chunk_size=500, overlap=50):
    chunks = []
    start = 0
    while start < len(text):
        end = start + chunk_size
        chunk = text[start:end]
        chunks.append(chunk)
        start = end - overlap
    return chunks

def get_embeddings(texts)-> np.ndarray:
    try:
        model = SentenceTransformer(emb_model, trust_remote_code=True)
        embeddings = model.encode(texts)

        print(f"Generated {len(embeddings)} embeddings.")
        return embeddings
    except Exception as e:
        st.error(f"Error generating embeddings: {e}")
        return []

def process_files(files):
    all_chunks = []
    all_embeddings = []
    chunks_metadata = []

    for file in files:
        print(f"Processing file: {file.name if hasattr(file, 'name') else os.path.basename(file)}")
        text = extract_text(file)
        if not text:  # Skip files that failed to process
            print(f"Skipping file {file.name if hasattr(file, 'name') else os.path.basename(file)} due to extraction error.")   
            continue
        print(f"Chunking text...{file.name if hasattr(file, 'name') else os.path.basename(file)}\n")
        chunks = chunk_text(text)
        embeddings = get_embeddings(chunks)
        # if not embeddings: # Skip files that failed to embed
        #     continue

        all_chunks.extend(chunks)
        all_embeddings.extend(embeddings)
        for i, chunk in enumerate(chunks):
            chunks_metadata.append({"file_name": file.name if hasattr(file, 'name') else os.path.basename(file), "chunk_index": i})
    return all_chunks, all_embeddings, chunks_metadata